"""Plain-text extraction for the file types a paper is usually submitted in."""
from __future__ import annotations

from pathlib import Path

SUPPORTED_SUFFIXES = {".txt", ".md", ".docx", ".pdf"}


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def _extract_docx(path: Path) -> str:
    import docx  # python-docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    # Blank lines between paragraphs so downstream paragraph-level analysis
    # sees the document's real paragraph structure.
    return "\n\n".join(p for p in parts if p.strip())


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)
